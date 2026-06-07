import re
import os
import pypdf
import docx

# Curated, comprehensive tech skills database
SKILLS_DB = [
    # Languages
    "python", "javascript", "typescript", "java", "c++", "c#", "go", "golang", "rust", 
    "ruby", "php", "swift", "kotlin", "scala", "r", "sql", "html", "css", "sass", "less",
    "bash", "shell", "powershell", "perl", "dart", "matlab",
    # Frontend Frameworks & Libraries
    "react", "react.js", "reactjs", "vue", "vue.js", "vuejs", "angular", "angularjs",
    "next.js", "nextjs", "nuxt", "svelte", "jquery", "tailwind", "tailwindcss",
    "bootstrap", "material ui", "mui", "redux", "redux-toolkit", "vuex", "pinia", "webpack", "vite",
    # Backend Frameworks
    "node.js", "nodejs", "express", "express.js", "django", "flask", "fastapi", 
    "spring", "spring boot", "laravel", "ruby on rails", "rails", "asp.net", "nest.js", "nestjs",
    # Cloud & DevOps
    "aws", "amazon web services", "gcp", "google cloud", "azure", "docker", "kubernetes", 
    "k8s", "terraform", "ansible", "jenkins", "gitlab ci", "github actions", "circleci", 
    "vagrant", "terraform", "prometheus", "grafana", "nginx", "apache", "linux", "unix",
    # Databases & Caching
    "postgresql", "postgres", "mysql", "mongodb", "sqlite", "redis", "elasticsearch", 
    "cassandra", "dynamodb", "mariadb", "oracle", "firebase", "supabase", "prisma", "sequelize",
    # Machine Learning & Data Science
    "machine learning", "ml", "deep learning", "dl", "artificial intelligence", "ai",
    "nlp", "natural language processing", "computer vision", "cv", "tensorflow", "pytorch", 
    "keras", "scikit-learn", "sklearn", "pandas", "numpy", "scipy", "seaborn", "matplotlib", 
    "tableau", "power bi", "powerbi", "hadoop", "spark", "apache spark", "jupyter", "opencv",
    # Mobile Development
    "ios", "android", "flutter", "react native", "xcode", "android studio", "swiftui",
    # Testing
    "selenium", "cypress", "playwright", "jest", "mocha", "junit", "testng", "postman",
    "unit testing", "integration testing", "automation testing",
    # Tools & Methodologies
    "git", "github", "gitlab", "jira", "confluence", "agile", "scrum", "kanban", "devops",
    "ci/cd", "rest api", "graphql", "grpc", "microservices", "system design", "docker-compose",
    # Design
    "figma", "adobe xd", "sketch", "photoshop", "illustrator", "ui/ux", "wireframing"
]

# Standardize skills name for clean presentation in UI
SKILL_DISPLAY_MAP = {
    "react.js": "React", "reactjs": "React", "react": "React",
    "vue.js": "Vue.js", "vuejs": "Vue.js", "vue": "Vue.js",
    "angularjs": "Angular", "angular": "Angular",
    "next.js": "Next.js", "nextjs": "Next.js",
    "tailwind": "Tailwind CSS", "tailwindcss": "Tailwind CSS",
    "node.js": "Node.js", "nodejs": "Node.js",
    "fastapi": "FastAPI", "django": "Django", "flask": "Flask",
    "spring boot": "Spring Boot", "spring": "Spring Boot",
    "aws": "AWS", "amazon web services": "AWS",
    "gcp": "GCP", "google cloud": "GCP",
    "azure": "Azure",
    "docker": "Docker", "kubernetes": "Kubernetes", "k8s": "Kubernetes",
    "terraform": "Terraform", "jenkins": "Jenkins",
    "postgresql": "PostgreSQL", "postgres": "PostgreSQL",
    "mysql": "MySQL", "mongodb": "MongoDB", "sqlite": "SQLite", "redis": "Redis",
    "python": "Python", "javascript": "JavaScript", "typescript": "TypeScript",
    "java": "Java", "c++": "C++", "c#": "C#", "golang": "Go", "go": "Go",
    "rust": "Rust", "swift": "Swift", "kotlin": "Kotlin", "dart": "Dart",
    "html": "HTML", "css": "CSS", "sql": "SQL",
    "machine learning": "Machine Learning", "ml": "Machine Learning",
    "deep learning": "Deep Learning", "dl": "Deep Learning",
    "nlp": "NLP", "tensorflow": "TensorFlow", "pytorch": "PyTorch",
    "scikit-learn": "Scikit-Learn", "sklearn": "Scikit-Learn",
    "pandas": "Pandas", "numpy": "NumPy",
    "git": "Git", "github": "GitHub", "jira": "Jira",
    "figma": "Figma", "ui/ux": "UI/UX", "rest api": "REST API",
    "graphql": "GraphQL", "microservices": "Microservices",
    "selenium": "Selenium", "cypress": "Cypress", "playwright": "Playwright"
}

def extract_text_from_pdf(pdf_path: str) -> str:
    """Extract text from a PDF file using pypdf"""
    text = ""
    try:
        with open(pdf_path, "rb") as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"Error reading PDF {pdf_path}: {e}")
    return text

def extract_text_from_docx(docx_path: str) -> str:
    """Extract text from a Word DOCX file"""
    text = ""
    try:
        doc = docx.Document(docx_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        print(f"Error reading DOCX {docx_path}: {e}")
    return text

def extract_text(file_path: str) -> str:
    """Generic text extractor handling PDF and DOCX"""
    _, ext = os.path.splitext(file_path.lower())
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        # Fallback to plain text read
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            print(f"Error reading file {file_path}: {e}")
            return ""

def clean_text(text: str) -> str:
    """Clean extracted text: strip double spaces, unwanted characters"""
    text = re.sub(r'\s+', ' ', text) # clean whitespace
    return text.strip()

def extract_email(text: str) -> str:
    """Extract email address from text using regex"""
    pattern = r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+'
    match = re.search(pattern, text)
    return match.group(0) if match else ""

def extract_phone(text: str) -> str:
    """Extract telephone number from text using standard regex pattern"""
    # Pattern to match numbers like +1-555-555-5555, (555) 555-5555, 555-555-5555, +91 9999999999
    pattern = r'(?:(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}|\+?\d{1,4}[-.\s]?\d{10})'
    match = re.search(pattern, text)
    return match.group(0) if match else ""

def extract_name(text: str, filename: str = "") -> str:
    """Heuristic name extraction: Check first few lines of text"""
    # Clean up empty lines and split
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    stop_words = {"resume", "curriculum", "vitae", "summary", "profile", "contact", "experience", "education", "skills", "about"}
    
    for line in lines[:5]: # Search first 5 lines
        # Check if line looks like a name: 2-4 words, starts with capital letters, no digits or special chars
        words = line.split()
        if 2 <= len(words) <= 4:
            # Check if any word is a common stop word or contains email/phone characters
            if any(w.lower() in stop_words for w in words):
                continue
            if "@" in line or any(c.isdigit() for c in line) or ":" in line:
                continue
            # Must look like proper capitalization
            if all(w[0].isupper() if w[0].isalpha() else False for w in words if w):
                return line
                
    # Fallback to filename if no name can be parsed
    if filename:
        name_part = os.path.splitext(os.path.basename(filename))[0]
        # Clean up separators like underscores or dashes
        name_part = re.sub(r'[-_]', ' ', name_part)
        name_part = re.sub(r'(?i)\bresume\b|\bcv\b', '', name_part).strip()
        if name_part:
            return name_part.title()
            
    return "Candidate Name"

def extract_skills(text: str) -> list:
    """Find all matching skills in resume text using word boundary searches"""
    text_lower = text.lower()
    found_skills = set()
    
    for skill in SKILLS_DB:
        # Avoid small skill strings from match mistakes (e.g. 'c' or 'r')
        if len(skill) <= 2:
            # Use strict word boundary for short skills
            pattern = r'\b' + re.escape(skill) + r'\b'
        else:
            # Normal check
            pattern = r'\b' + re.escape(skill) + r'\b'
            
        # Add boundary for special characters in skill (e.g. c++, c#)
        if skill in ["c++", "c#"]:
            pattern = r'(?:\b|(?<=\s))' + re.escape(skill) + r'(?:\b|(?=\s))'
            
        if re.search(pattern, text_lower):
            display_name = SKILL_DISPLAY_MAP.get(skill, skill.title())
            found_skills.add(display_name)
            
    return sorted(list(found_skills))

def parse_resume(file_path: str, filename: str = "") -> dict:
    """Parse a resume file and return a dictionary of properties"""
    raw_text = extract_text(file_path)
    cleaned_text = clean_text(raw_text)
    
    email = extract_email(cleaned_text)
    phone = extract_phone(cleaned_text)
    name = extract_name(raw_text, filename or file_path)
    skills = extract_skills(cleaned_text)
    
    return {
        "name": name,
        "email": email,
        "phone": phone,
        "skills": skills,
        "raw_text": cleaned_text
    }

def parse_resume_content(raw_text: str) -> dict:
    """
    Parse raw pasted text (from ChatGPT/Claude/Gemini etc.) into structured
    resume sections. Returns a dict with all the sections needed for template rendering.
    """
    lines = [l.strip() for l in raw_text.split('\n')]
    text = raw_text

    # ── Contact Info ──────────────────────────────────────
    email = extract_email(text)
    phone = extract_phone(text)

    # LinkedIn
    linkedin_match = re.search(r'linkedin\.com/in/[\w\-]+', text, re.IGNORECASE)
    linkedin = linkedin_match.group(0) if linkedin_match else ''

    # GitHub
    github_match = re.search(r'github\.com/[\w\-]+', text, re.IGNORECASE)
    github = github_match.group(0) if github_match else ''

    # Portfolio / Website
    website_match = re.search(r'https?://(?!linkedin|github)[\w.\-/]+', text, re.IGNORECASE)
    website = website_match.group(0) if website_match else ''

    # ── Name ──────────────────────────────────────────────
    name = extract_name(text)

    # ── Job Title / Tagline (line right after name) ───────
    title = ''
    non_empty = [l for l in lines if l.strip()]
    if non_empty:
        # First line that's not the name and not contact info
        for ln in non_empty[:6]:
            if ln == name:
                continue
            if re.search(r'@|\d{7,}|linkedin|github|http', ln, re.IGNORECASE):
                continue
            if 2 <= len(ln.split()) <= 8 and not ln.endswith(':'):
                title = ln
                break

    # ── Section headers detection ─────────────────────────
    SECTION_PATTERNS = {
        'summary':        r'^(summary|profile|about|objective|about me|professional summary|career objective)\s*:?\s*$',
        'experience':     r'^(experience|work experience|professional experience|employment|employment history|work history|career history)\s*:?\s*$',
        'education':      r'^(education|academic background|academic qualifications|qualifications)\s*:?\s*$',
        'skills':         r'^(skills|technical skills|core competencies|competencies|technologies|tech stack|key skills|skill set)\s*:?\s*$',
        'projects':       r'^(projects|personal projects|key projects|notable projects|portfolio)\s*:?\s*$',
        'certifications': r'^(certifications?|certificates?|licenses?|credentials?|achievements?|awards?)\s*:?\s*$',
        'languages':      r'^(languages?|spoken languages?)\s*:?\s*$',
        'interests':      r'^(interests?|hobbies|extracurricular)\s*:?\s*$',
    }

    sections = {}
    current_section = None
    current_content = []

    for line in lines:
        matched = False
        for section, pattern in SECTION_PATTERNS.items():
            if re.match(pattern, line.strip(), re.IGNORECASE):
                if current_section and current_content:
                    sections[current_section] = '\n'.join(current_content).strip()
                current_section = section
                current_content = []
                matched = True
                break
        if not matched and current_section:
            current_content.append(line)

    if current_section and current_content:
        sections[current_section] = '\n'.join(current_content).strip()

    # ── Parse Experience entries ───────────────────────────
    def parse_experience(text_block: str) -> list:
        if not text_block:
            return []
        entries = []
        # Split on blank lines or patterns that start a new entry
        blocks = re.split(r'\n\s*\n', text_block)
        for block in blocks:
            block = block.strip()
            if not block:
                continue
            block_lines = [l.strip() for l in block.split('\n') if l.strip()]
            if not block_lines:
                continue
            entry = {
                'title': '',
                'company': '',
                'dates': '',
                'location': '',
                'bullets': []
            }
            # First line: role title
            entry['title'] = block_lines[0]
            # Second line: company + dates + location
            if len(block_lines) > 1:
                second = block_lines[1]
                # Extract dates
                dates_match = re.search(
                    r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[\w\s,]*\d{4}\s*[-–—]\s*(?:Present|Current|Now|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[\w\s,]*\d{4})|\d{4}\s*[-–—]\s*(?:Present|\d{4}))',
                    second, re.IGNORECASE
                )
                if dates_match:
                    entry['dates'] = dates_match.group(0).strip()
                    second = second.replace(dates_match.group(0), '').strip().strip('|,·•-').strip()
                # Remaining: company | location
                parts = re.split(r'\s*[|,·•]\s*', second)
                if parts:
                    entry['company'] = parts[0].strip()
                if len(parts) > 1:
                    entry['location'] = parts[1].strip()
            # Rest: bullet points
            for bl in block_lines[2:]:
                bl = bl.lstrip('•*-·–—').strip()
                if bl:
                    entry['bullets'].append(bl)
            entries.append(entry)
        return entries

    def parse_education(text_block: str) -> list:
        if not text_block:
            return []
        entries = []
        blocks = re.split(r'\n\s*\n', text_block)
        for block in blocks:
            block = block.strip()
            if not block:
                continue
            block_lines = [l.strip() for l in block.split('\n') if l.strip()]
            if not block_lines:
                continue
            entry = {'degree': '', 'school': '', 'dates': '', 'gpa': '', 'details': []}
            entry['degree'] = block_lines[0]
            if len(block_lines) > 1:
                second = block_lines[1]
                dates_match = re.search(r'(\d{4}\s*[-–—]\s*(?:Present|\d{4})|\d{4})', second)
                if dates_match:
                    entry['dates'] = dates_match.group(0).strip()
                    second = second.replace(dates_match.group(0), '').strip().strip('|,·•-').strip()
                gpa_match = re.search(r'GPA\s*[:.]?\s*([\d.]+)', block_lines[-1], re.IGNORECASE)
                if gpa_match:
                    entry['gpa'] = gpa_match.group(1)
                entry['school'] = second
            for bl in block_lines[2:]:
                bl = bl.lstrip('•*-·').strip()
                if bl and 'GPA' not in bl:
                    entry['details'].append(bl)
            entries.append(entry)
        return entries

    def parse_skills(text_block: str) -> list:
        """Return list of skill strings"""
        if not text_block:
            return []
        # Remove category labels like "Languages:" etc.
        text_block = re.sub(r'^[\w\s]+:\s*', '', text_block, flags=re.MULTILINE)
        # Split on commas, bullets, newlines, semicolons
        raw = re.split(r'[,\n•|;/]+', text_block)
        skills = []
        for s in raw:
            s = s.strip().lstrip('*-·').strip()
            if s and len(s) < 50:
                skills.append(s)
        return [s for s in skills if s]

    def parse_projects(text_block: str) -> list:
        if not text_block:
            return []
        entries = []
        blocks = re.split(r'\n\s*\n', text_block)
        for block in blocks:
            block = block.strip()
            if not block:
                continue
            block_lines = [l.strip() for l in block.split('\n') if l.strip()]
            if not block_lines:
                continue
            entry = {'name': block_lines[0], 'tech': '', 'bullets': []}
            tech_match = re.search(r'Tech(?:nologies|nology|nical stack)?[:\s]+([^\n]+)', block, re.IGNORECASE)
            if tech_match:
                entry['tech'] = tech_match.group(1).strip()
            for bl in block_lines[1:]:
                bl = bl.lstrip('•*-·').strip()
                if bl and not bl.startswith('Tech'):
                    entry['bullets'].append(bl)
            entries.append(entry)
        return entries

    def parse_certifications(text_block: str) -> list:
        if not text_block:
            return []
        items = []
        for line in text_block.split('\n'):
            line = line.lstrip('•*-·').strip()
            if line:
                items.append(line)
        return items

    # ── Location ──────────────────────────────────────────
    location_match = re.search(
        r'\b([A-Z][a-z]+(?:\s[A-Z][a-z]+)?,\s*(?:[A-Z]{2,3}|[A-Z][a-z]+))\b',
        text
    )
    location = location_match.group(0) if location_match else ''

    return {
        'name': name,
        'title': title,
        'email': email,
        'phone': phone,
        'linkedin': linkedin,
        'github': github,
        'website': website,
        'location': location,
        'summary': sections.get('summary', '').strip(),
        'skills': parse_skills(sections.get('skills', '')),
        'experience': parse_experience(sections.get('experience', '')),
        'education': parse_education(sections.get('education', '')),
        'projects': parse_projects(sections.get('projects', '')),
        'certifications': parse_certifications(sections.get('certifications', '')),
        'languages': parse_skills(sections.get('languages', '')),
        'interests': parse_certifications(sections.get('interests', '')),
    }


if __name__ == "__main__":
    # Test script locally
    test_resume = """
    John Doe
    Software Engineer
    john.doe@example.com | (555) 123-4567
    
    SUMMARY
    Highly motivated developer with experience building apps.
    
    SKILLS
    React, Python, Node.js, AWS, SQL, and Docker. 
    Also familiar with Machine Learning and TensorFlow.
    """
    
    print("Testing parser with mockup text:")
    email = extract_email(test_resume)
    phone = extract_phone(test_resume)
    name = extract_name(test_resume)
    skills = extract_skills(test_resume)
    
    print(f"Name: {name}")
    print(f"Email: {email}")
    print(f"Phone: {phone}")
    print(f"Skills: {skills}")

